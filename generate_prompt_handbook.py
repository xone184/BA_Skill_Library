import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

BASE_DIR = r"e:\BA Skill Library\BA-Skills"
OUTPUT_FILE = r"e:\BA Skill Library\BA_Skill_Prompt_Handbook.docx"

# ================================================================
# BẢNG PROMPT MẪU CHUYÊN SÂU theo từng nhóm
# (Được soạn thủ công kết hợp với nội dung từ các file .prompt)
# ================================================================
PROMPT_TEMPLATES = [
    # ============================================================
    # NHÓM 1: PHÂN TÍCH NGHIỆP VỤ DOMAIN
    # ============================================================
    {
        "group": "01 · Phân tích Nghiệp vụ Domain",
        "skills": [
            {
                "skill": "phan-tich-lead-management",
                "name": "Phân tích Lead Management (CRM)",
                "use_case": "Phân tích quy trình quản lý khách hàng tiềm năng từ đầu vào đến chuyển đổi",
                "prompt": (
                    "/phan-tich-lead-management\n"
                    "Phân tích quy trình Lead Management cho hệ thống CRM sau:\n"
                    "- Nguồn Lead: [Web form / Zalo / Call center / ...]\n"
                    "- Quy trình xử lý: [Mô tả ngắn]\n"
                    "- Điều kiện chuyển trạng thái: [New → Contacted → Qualified → Converted]\n\n"
                    "Yêu cầu:\n"
                    "1. Phân tích đầy đủ vòng đời Lead (Lead Lifecycle)\n"
                    "2. Xác định Business Rules cho từng bước chuyển trạng thái\n"
                    "3. Đề xuất các trường dữ liệu bắt buộc và tùy chọn\n"
                    "4. Nêu rõ phân quyền: Sales / Manager / System thực hiện gì"
                ),
                "note": "Thay phần trong [] bằng thông tin thực tế của dự án."
            },
            {
                "skill": "phan-tich-opportunity",
                "name": "Phân tích Opportunity (CRM)",
                "use_case": "Phân tích quy trình quản lý cơ hội bán hàng (Sales Pipeline)",
                "prompt": (
                    "/phan-tich-opportunity\n"
                    "Phân tích Opportunity Management cho dự án [Tên dự án]:\n"
                    "- Các giai đoạn Sales Pipeline: [Qualify → Propose → Negotiate → Won/Lost]\n"
                    "- Điều kiện thắng/thua cơ hội: [Mô tả]\n\n"
                    "Yêu cầu:\n"
                    "1. Liệt kê đầy đủ Functional Requirements (FR-OPP-001...)\n"
                    "2. Xây dựng bảng Business Rules quyết định Won/Lost\n"
                    "3. Vẽ State Machine Diagram vòng đời Opportunity"
                ),
                "note": ""
            },
            {
                "skill": "phan-tich-production-order",
                "name": "Phân tích Lệnh Sản xuất (MES)",
                "use_case": "Phân tích quy trình tạo và thực thi lệnh sản xuất trong nhà máy",
                "prompt": (
                    "/phan-tich-production-order\n"
                    "Phân tích quy trình Lệnh Sản xuất cho nhà máy [Tên nhà máy / Ngành nghề]:\n"
                    "- Quy trình: Lập kế hoạch → Phát lệnh → Thực thi tại máy → Báo cáo hoàn thành\n"
                    "- Loại máy/thiết bị: [CNC / Phun nhựa / Đóng gói / ...]\n"
                    "- Hệ thống liên quan: ERP, MES, WMS\n\n"
                    "Yêu cầu:\n"
                    "1. Phân tích luồng nghiệp vụ chi tiết (Business Process Flow)\n"
                    "2. Xác định trạng thái lệnh sản xuất (State Machine)\n"
                    "3. Liệt kê các Business Rules về vật tư, định mức, và kiểm tra chất lượng\n"
                    "4. Đề xuất bảng dữ liệu và các trường cốt lõi"
                ),
                "note": "Skill chuyên biệt MES. Kết hợp với /phan-tich-oee để phân tích hiệu suất máy."
            },
            {
                "skill": "phan-tich-inventory-management",
                "name": "Phân tích Quản lý Tồn kho (WMS)",
                "use_case": "Phân tích nghiệp vụ nhập - xuất - tồn kho trong kho hàng",
                "prompt": (
                    "/phan-tich-inventory-management\n"
                    "Phân tích nghiệp vụ Quản lý Tồn kho cho kho [Loại kho: Thương phẩm / Lạnh / ...]:\n"
                    "- Loại nghiệp vụ cần phân tích: [Nhập kho / Xuất kho / Điều chỉnh / Kiểm kê]\n"
                    "- Đơn vị tính: [Thùng, Pallet, Kg, ...]\n"
                    "- Phương pháp xuất kho: [FIFO / FEFO / LIFO]\n\n"
                    "Yêu cầu:\n"
                    "1. Phân tích đầy đủ luồng nhập-xuất-tồn\n"
                    "2. Xác định các quy tắc kiểm soát: Min/Max, Safety Stock, Reorder Point\n"
                    "3. Đề xuất cấu trúc dữ liệu kho và phân vùng (Location / Zone)"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 2: PHÂN TÍCH YÊU CẦU
    # ============================================================
    {
        "group": "02 · Phân tích & Viết Tài liệu Yêu cầu",
        "skills": [
            {
                "skill": "write-srs-ultra",
                "name": "Viết SRS (Ultra)",
                "use_case": "Tổng hợp tất cả thông tin đã phân tích thành tài liệu SRS hoàn chỉnh",
                "prompt": (
                    "/write-srs-ultra\n"
                    "Hãy tổng hợp TẤT CẢ các thông tin chúng ta vừa thảo luận "
                    "(bao gồm Use Case, Business Rules, State Machine, ERD, và Wireframe) "
                    "thành một bản tài liệu SRS hoàn chỉnh cho module [Tên module].\n\n"
                    "Cấu trúc SRS cần có:\n"
                    "1. Tổng quan hệ thống (System Overview)\n"
                    "2. Danh mục Actors và Use Cases\n"
                    "3. Functional Requirements (FR) - Đánh mã FR-[MODULE]-001\n"
                    "4. Non-Functional Requirements (NFR)\n"
                    "5. Business Rules (BR) - Đánh mã BR-[MODULE]-001\n"
                    "6. Data Dictionary (Mô tả các trường dữ liệu quan trọng)\n"
                    "7. Acceptance Criteria (Given-When-Then)"
                ),
                "note": "Dùng sau khi đã chạy xong các skill phân tích nhỏ hơn. Đây là 'Cú chốt' của Skill Chaining."
            },
            {
                "skill": "phan-tich-use-case",
                "name": "Phân tích Use Case",
                "use_case": "Đặc tả chi tiết Use Case Specification bao gồm cả các luồng lỗi và ngoại lệ",
                "prompt": (
                    "/phan-tich-use-case\n"
                    "Đặc tả chi tiết Use Case cho tính năng: [Tên tính năng]\n"
                    "- Actor chính: [Người dùng / Nhân viên / Quản lý]\n"
                    "- Mô tả tóm tắt: [1-2 câu mô tả tính năng]\n\n"
                    "Yêu cầu:\n"
                    "1. Viết đầy đủ Use Case Specification bao gồm:\n"
                    "   - Basic Flow (Happy Path)\n"
                    "   - Alternate Flows (Luồng rẽ nhánh hợp lệ)\n"
                    "   - Exception Flows (Luồng lỗi)\n"
                    "2. Ghi rõ Pre-conditions và Post-conditions\n"
                    "3. [Micro-tasking] Chỉ viết Exception Flows cho bước [Số bước]"
                ),
                "note": "Thêm câu [Micro-tasking] cuối cùng khi chỉ muốn AI xử lý 1 phần nhỏ."
            },
            {
                "skill": "write-user-story",
                "name": "Viết User Story",
                "use_case": "Viết User Stories chuẩn Agile kèm Acceptance Criteria theo Gherkin",
                "prompt": (
                    "/write-user-story\n"
                    "Viết User Stories cho tính năng [Tên tính năng] dựa trên mô tả sau:\n"
                    "[Mô tả nghiệp vụ hoặc dán đoạn SRS / Use Case vào đây]\n\n"
                    "Yêu cầu:\n"
                    "1. Chuẩn format: 'Là [vai trò], tôi muốn [mục tiêu] để [lợi ích]'\n"
                    "2. Phân loại độ ưu tiên MoSCoW (Must/Should/Could/Won't)\n"
                    "3. Viết Acceptance Criteria dạng Gherkin (Given-When-Then) cho từng Story\n"
                    "4. Phải bao gồm cả Happy Path và Exception Flow trong AC"
                ),
                "note": ""
            },
            {
                "skill": "phan-tich-business-rules",
                "name": "Phân tích Business Rules",
                "use_case": "Bóc tách và chuẩn hóa toàn bộ quy tắc nghiệp vụ từ mô tả thô",
                "prompt": (
                    "/phan-tich-business-rules\n"
                    "Từ mô tả nghiệp vụ sau, hãy bóc tách toàn bộ Business Rules:\n"
                    "[Dán mô tả nghiệp vụ / đoạn SRS / biên bản họp vào đây]\n\n"
                    "Yêu cầu:\n"
                    "1. Phân loại theo nhóm: Validation Rules / Calculation Rules / Workflow Rules\n"
                    "2. Đánh mã: BR-[MODULE]-001, BR-[MODULE]-002...\n"
                    "3. Với mỗi rule phức tạp, tạo Decision Table để làm rõ\n"
                    "4. Đề xuất Test Case tương ứng cho từng Business Rule quan trọng"
                ),
                "note": "Đặc biệt hữu ích khi có nhiều điều kiện AND/OR lồng nhau."
            },
            {
                "skill": "review-requirement",
                "name": "Review Yêu cầu",
                "use_case": "Kiểm tra chất lượng tài liệu yêu cầu theo chuẩn CLEAR, COMPLETE, CONSISTENT",
                "prompt": (
                    "/review-requirement\n"
                    "Hãy review tài liệu yêu cầu sau theo tiêu chuẩn chất lượng BA:\n"
                    "[Dán nội dung SRS / User Story / Use Case cần review]\n\n"
                    "Tiêu chí đánh giá:\n"
                    "- CLEAR: Có rõ ràng và không mơ hồ không?\n"
                    "- COMPLETE: Có thiếu luồng nào không (đặc biệt Exception Flows)?\n"
                    "- CONSISTENT: Có mâu thuẫn giữa các yêu cầu không?\n"
                    "- TESTABLE: Có thể viết Test Case từ yêu cầu này không?\n"
                    "- TRACEABLE: Mỗi yêu cầu có ID và nguồn gốc rõ ràng không?\n\n"
                    "Xuất kết quả dạng bảng: [Mục] | [Vấn đề] | [Gợi ý sửa]"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 3: VẼ SƠ ĐỒ & MÔ HÌNH HÓA
    # ============================================================
    {
        "group": "03 · Vẽ Sơ đồ & Mô hình hóa",
        "skills": [
            {
                "skill": "activity-diagram",
                "name": "Vẽ Activity Diagram (Lưu đồ dọc - Chuẩn Visio)",
                "use_case": "Vẽ lưu đồ phân làn dọc (Vertical Swimlanes) kiểu Visio/Draw.io bằng PlantUML",
                "prompt": (
                    "/activity-diagram\n"
                    "Từ mô tả quy trình (hoặc SRS) sau, hãy vẽ Activity Diagram bằng mã PlantUML:\n"
                    "[Dán mô tả quy trình hoặc các bước nghiệp vụ]\n\n"
                    "YÊU CẦU BẮT BUỘC:\n"
                    "1. Dùng PlantUML với cú pháp Vertical Swimlanes (|Người dùng| và |Hệ thống|)\n"
                    "2. Luồng đi từ trên xuống dưới (Top-Down)\n"
                    "3. Chỉ có 1 điểm Bắt đầu (start) và tối đa 2 điểm Kết thúc (stop)\n"
                    "4. Mọi nhánh lỗi / hủy phải hội tụ về End Node, không thả nổi\n"
                    "5. Tên mỗi bước phải là: Động từ + Tân ngữ (VD: Kiểm tra điều kiện)"
                ),
                "note": "Kết quả mã PlantUML: Mở Draw.io → Insert → Advanced → PlantUML → Dán vào."
            },
            {
                "skill": "activity-diagram",
                "name": "Vẽ Activity Diagram (Lưu đồ ngang - Mermaid)",
                "use_case": "Vẽ lưu đồ phân làn ngang (Horizontal Swimlanes) bằng Mermaid.js để xem trực tiếp trong Claude/AI",
                "prompt": (
                    "/activity-diagram\n"
                    "Từ mô tả quy trình sau, hãy vẽ Activity Diagram bằng mã Mermaid:\n"
                    "[Dán mô tả quy trình]\n\n"
                    "YÊU CẦU BẮT BUỘC:\n"
                    "1. Dùng cú pháp Mermaid `flowchart LR` (Luồng từ Trái sang Phải)\n"
                    "2. Swimlane nằm ngang: 1 subgraph cho Người dùng, 1 subgraph cho Hệ thống\n"
                    "3. Gom nhóm hành động cùng Actor trước khi chuyển làn (tránh đường cắt chéo)\n"
                    "4. Quy tụ mọi nhánh phụ (Hủy, Lỗi) về 1 điểm Kết thúc (End) duy nhất ở cuối\n"
                    "5. Không dùng ký tự đặc biệt trong ID node"
                ),
                "note": "Mermaid tự render trực tiếp trên Claude. Không cần export ra file ngoài."
            },
            {
                "skill": "draw-bpmn",
                "name": "Vẽ BPMN",
                "use_case": "Vẽ sơ đồ BPMN chuẩn với Pool, Lanes và Gateway cho quy trình nhiều phòng ban",
                "prompt": (
                    "/draw-bpmn\n"
                    "Từ SRS sau, hãy viết mã PlantUML (hoặc Mermaid) BPMN để tôi dán vào Draw.io:\n"
                    "[Dán mô tả quy trình / SRS]\n\n"
                    "YÊU CẦU BẮT BUỘC:\n"
                    "1. Pool = Hệ thống [Tên hệ thống]. Lanes = Vai trò (VD: Nhân viên KD / Kế toán / Hệ thống)\n"
                    "2. Bắt buộc dùng layout trái → phải (`left to right direction` hoặc `flowchart LR`)\n"
                    "3. Phân biệt rõ: User Task (Hành động người dùng) và Service Task (Hành động tự động)\n"
                    "4. Mọi Gateway tách nhánh (XOR/AND) phải có Gateway tương ứng để gom lại (Merge)\n"
                    "5. Chỉ có 1 Start Event và tối đa 2 End Events (Thành công / Thất bại)"
                ),
                "note": "Dùng cho quy trình liên phòng ban. Trong MES: dùng BPMN cho quy trình xuyên bộ phận, Activity Diagram cho thao tác tại 1 trạm máy."
            },
            {
                "skill": "thiet-ke-erd",
                "name": "Thiết kế ERD (Entity Relationship Diagram)",
                "use_case": "Thiết kế cấu trúc cơ sở dữ liệu và mối quan hệ giữa các bảng",
                "prompt": (
                    "/thiet-ke-erd\n"
                    "Thiết kế ERD cho module [Tên module] với các thực thể sau:\n"
                    "[Liệt kê các đối tượng nghiệp vụ: VD: Đơn hàng, Khách hàng, Sản phẩm, ...]\n\n"
                    "Yêu cầu:\n"
                    "1. Vẽ ERD dạng Crow's Foot (thể hiện đầy đủ cardinality: 1-1, 1-N, N-N)\n"
                    "2. Đặt tên bảng theo chuẩn: snake_case, số nhiều (VD: work_orders, production_lines)\n"
                    "3. Primary Key: [tên_bảng_số_ít]_id. Foreign Key rõ ràng\n"
                    "4. Tối thiểu đạt chuẩn 3NF (Third Normal Form)\n"
                    "5. Ghi chú Đơn vị đo lường (UoM) cho các trường số lượng (đặc biệt quan trọng với MES)"
                ),
                "note": ""
            },
            {
                "skill": "ve-sequence-diagram",
                "name": "Vẽ Sequence Diagram",
                "use_case": "Mô tả luồng giao tiếp chi tiết giữa các component/service trong hệ thống",
                "prompt": (
                    "/ve-sequence-diagram\n"
                    "Vẽ Sequence Diagram cho luồng [Tên luồng, VD: Đặt hàng qua API]:\n"
                    "- Các Actor / Component tham gia: [Frontend, Backend API, Database, 3rd Party]\n"
                    "- Luồng chính: [Mô tả từng bước]\n\n"
                    "Yêu cầu:\n"
                    "1. Dùng PlantUML `@startuml ... @enduml`\n"
                    "2. Mọi message (mũi tên) phải có nhãn rõ ràng (tên hàm hoặc endpoint)\n"
                    "3. Dùng combined fragments: `alt` cho rẽ nhánh, `loop` cho vòng lặp, `opt` cho tùy chọn\n"
                    "4. Thể hiện cả luồng lỗi (Exception): Dùng `alt` với nhánh [Error]"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 4: THIẾT KẾ UI/UX
    # ============================================================
    {
        "group": "04 · Thiết kế Giao diện (UI/UX)",
        "skills": [
            {
                "skill": "phan-tich-man-hinh-crud",
                "name": "Phân tích màn hình CRUD",
                "use_case": "Đặc tả chi tiết màn hình danh sách, thêm mới, sửa, xóa cho 1 đối tượng",
                "prompt": (
                    "/phan-tich-man-hinh-crud\n"
                    "Phân tích và đặc tả màn hình CRUD cho đối tượng: [Tên đối tượng, VD: Lệnh sản xuất]\n"
                    "- Vai trò có quyền truy cập: [Admin / Manager / Operator]\n"
                    "- Các trường dữ liệu hiển thị: [Liệt kê]\n\n"
                    "Yêu cầu:\n"
                    "1. Đặc tả 4 màn hình: Danh sách (List) / Thêm mới (Create) / Sửa (Edit) / Xem chi tiết (Detail)\n"
                    "2. Mỗi màn hình phải có Screen ID (VD: SCR-WO-001)\n"
                    "3. Thể hiện đủ 5 trạng thái: Default / Empty / Loading / Error / Success\n"
                    "4. Đặc tả rõ các nút bấm (Button), validation message và quyền hiển thị theo Role"
                ),
                "note": ""
            },
            {
                "skill": "phan-tich-mobile-app",
                "name": "Phân tích giao diện Mobile/Tablet (MES)",
                "use_case": "Thiết kế giao diện tối ưu cho Tablet/Máy quét tại xưởng sản xuất",
                "prompt": (
                    "/phan-tich-mobile-app\n"
                    "Thiết kế giao diện Tablet/Mobile cho màn hình [Tên màn hình] tại xưởng sản xuất:\n"
                    "- Thiết bị: [Tablet 10 inch / Máy quét barcode / ...]\n"
                    "- Người dùng: [Công nhân / Quản đốc]\n"
                    "- Tính năng chính: [Xác nhận sản lượng / Quét mã vật tư / ...]\n\n"
                    "Yêu cầu:\n"
                    "1. Tối ưu cho thao tác bằng ngón tay (Nút bấm tối thiểu 44x44px)\n"
                    "2. Font chữ lớn, tương phản cao (dùng được ngoài ánh sáng xưởng)\n"
                    "3. Hỗ trợ chế độ Offline: Lưu Queue nội bộ khi mất kết nối WIFI, đồng bộ khi có lại\n"
                    "4. Phản hồi âm thanh/rung cho môi trường ồn ào"
                ),
                "note": "Chuyên biệt cho MES. Kết hợp với /thiet-ke-role-based-ui cho màn hình Web quản lý."
            },
            {
                "skill": "thiet-ke-role-based-ui",
                "name": "Thiết kế UI theo Role (Phân quyền)",
                "use_case": "Thiết kế giao diện khác nhau cho từng vai trò người dùng",
                "prompt": (
                    "/thiet-ke-role-based-ui\n"
                    "Thiết kế giao diện phân quyền theo Role cho module [Tên module]:\n"
                    "- Danh sách Roles: [Admin / Manager / Operator / Viewer]\n"
                    "- Đặc điểm từng Role: [Mô tả quyền hạn]\n\n"
                    "Yêu cầu:\n"
                    "1. Vẽ ma trận phân quyền: Role × Chức năng → (Xem / Sửa / Xóa / Ẩn)\n"
                    "2. Mô tả cụ thể UI thay đổi thế nào theo từng Role (VD: Button Xóa ẩn với Viewer)\n"
                    "3. Xử lý Navigation Menu theo Role\n"
                    "4. Đặc tả hành vi khi người dùng truy cập URL không có quyền"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 5: KIỂM THỬ (TESTING)
    # ============================================================
    {
        "group": "05 · Kiểm thử (Testing & QA)",
        "skills": [
            {
                "skill": "viet-test-case-ultra",
                "name": "Viết Test Case (Ultra)",
                "use_case": "Sinh Test Case chuyên sâu bằng các kỹ thuật BVA, Equivalence Partitioning, Decision Table",
                "prompt": (
                    "/viet-test-case-ultra\n"
                    "Viết Test Case cho tính năng [Tên tính năng] với kỹ thuật [Chọn 1]:\n"
                    "  A) Boundary Value Analysis (BVA) - Phân tích giá trị biên\n"
                    "  B) Equivalence Partitioning - Phân vùng tương đương\n"
                    "  C) Decision Table - Bảng quyết định\n\n"
                    "Thông tin đầu vào:\n"
                    "- Business Rule cần test: [VD: Tuổi phải từ 18 đến 60]\n"
                    "- Dữ liệu hiện tại: [Mô tả]\n\n"
                    "[Micro-tasking nếu cần] Chỉ sinh Test Case BVA cho trường [Tên trường]."
                ),
                "note": "BVA tự động sinh đủ 6 cases: Min-1, Min, Min+1, Max-1, Max, Max+1."
            },
            {
                "skill": "ke-hoach-uat",
                "name": "Lập Kế hoạch UAT",
                "use_case": "Lập kế hoạch kiểm thử chấp nhận người dùng (User Acceptance Testing)",
                "prompt": (
                    "/ke-hoach-uat\n"
                    "Lập Kế hoạch UAT cho dự án [Tên dự án / Module]:\n"
                    "- Phạm vi kiểm thử: [Danh sách tính năng cần UAT]\n"
                    "- Thời gian: [Ngày bắt đầu → Ngày kết thúc]\n"
                    "- Nhân sự UAT: [Danh sách người tham gia và vai trò]\n\n"
                    "Yêu cầu:\n"
                    "1. Kế hoạch chi tiết theo ngày (Daily Plan)\n"
                    "2. Danh sách Kịch bản UAT (Test Scenarios) ưu tiên theo nghiệp vụ quan trọng\n"
                    "3. Quy trình báo cáo lỗi (Bug Report Process) và điều kiện PASS/FAIL\n"
                    "4. Entry/Exit Criteria cho toàn bộ đợt UAT"
                ),
                "note": ""
            },
            {
                "skill": "viet-kich-ban-uat-nang-cao",
                "name": "Viết Kịch bản UAT Nâng cao",
                "use_case": "Viết kịch bản UAT thực tế, bao gồm cả kịch bản nghiệp vụ đầu-cuối (End-to-End)",
                "prompt": (
                    "/viet-kich-ban-uat-nang-cao\n"
                    "Viết Kịch bản UAT cho luồng nghiệp vụ đầu-cuối (E2E): [Mô tả luồng]\n"
                    "VD: Từ lúc Lead được tạo → Qualified → Tạo Opportunity → Won → Tạo hợp đồng\n\n"
                    "Yêu cầu:\n"
                    "1. Kịch bản Happy Path (Mọi thứ diễn ra bình thường)\n"
                    "2. Kịch bản Negative (Nhập sai / Thiếu dữ liệu / Vượt ngưỡng giới hạn)\n"
                    "3. Kịch bản Regression (Tính năng cũ không bị ảnh hưởng sau khi thêm tính năng mới)\n"
                    "4. Format: [Kịch bản] | [Bước thực hiện] | [Kết quả mong đợi] | [Kết quả thực tế] | [Pass/Fail]"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 6: DỮ LIỆU & TÍCH HỢP
    # ============================================================
    {
        "group": "06 · Dữ liệu & Tích hợp API",
        "skills": [
            {
                "skill": "thiet-ke-api-contract",
                "name": "Thiết kế API Contract",
                "use_case": "Đặc tả chi tiết API Contract (Request/Response/Error) theo chuẩn OpenAPI",
                "prompt": (
                    "/thiet-ke-api-contract\n"
                    "Thiết kế API Contract cho [Tên API / Tính năng] theo chuẩn RESTful:\n"
                    "- Endpoint cần thiết kế: [VD: POST /api/v1/work-orders]\n"
                    "- Mục đích: [Mô tả]\n"
                    "- Hệ thống gọi (Caller): [Frontend / Mobile App / ERP]\n\n"
                    "Yêu cầu:\n"
                    "1. Method, URL, Headers (Authorization)\n"
                    "2. Request Body (JSON Schema + Validation Rules)\n"
                    "3. Response thành công (200/201) và Response lỗi (400/401/404/500)\n"
                    "4. Ví dụ cụ thể (Example) cho Request và Response\n"
                    "5. Ghi chú Business Rules liên quan đến API này"
                ),
                "note": ""
            },
            {
                "skill": "create-data-dictionary",
                "name": "Tạo Data Dictionary",
                "use_case": "Tạo từ điển dữ liệu mô tả chi tiết tất cả các trường của hệ thống",
                "prompt": (
                    "/create-data-dictionary\n"
                    "Tạo Data Dictionary cho bảng / module [Tên bảng / module]:\n"
                    "- Danh sách các trường: [Liệt kê hoặc dán ERD vào]\n\n"
                    "Yêu cầu format:\n"
                    "| Tên trường | Kiểu dữ liệu | Bắt buộc | Giá trị mặc định | Mô tả | Validation Rule | Ví dụ |\n\n"
                    "Lưu ý:\n"
                    "- Ghi rõ Đơn vị đo lường (UoM) cho trường số lượng (đặc biệt quan trọng với MES)\n"
                    "- Ghi rõ Enum values cho các trường trạng thái (Status)\n"
                    "- Đánh dấu [FK] cho Foreign Key và bảng tham chiếu"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 7: PHÂN TÍCH & CHIẾN LƯỢC
    # ============================================================
    {
        "group": "07 · Phân tích Chiến lược & Dự án",
        "skills": [
            {
                "skill": "gap-analysis",
                "name": "Phân tích GAP",
                "use_case": "So sánh quy trình hiện tại (As-Is) với quy trình mong muốn (To-Be) để tìm khoảng cách",
                "prompt": (
                    "/gap-analysis\n"
                    "Phân tích GAP cho quy trình [Tên quy trình]:\n\n"
                    "AS-IS (Hiện tại):\n"
                    "[Mô tả quy trình đang làm, các vấn đề, điểm đau]\n\n"
                    "TO-BE (Mong muốn):\n"
                    "[Mô tả quy trình lý tưởng, mục tiêu cần đạt]\n\n"
                    "Yêu cầu:\n"
                    "1. Bảng GAP Analysis: [Lĩnh vực] | [As-Is] | [To-Be] | [Khoảng cách] | [Giải pháp]\n"
                    "2. Đánh giá mức độ ưu tiên (Cao / Trung / Thấp) cho từng khoảng cách\n"
                    "3. Đề xuất lộ trình triển khai (Roadmap)"
                ),
                "note": ""
            },
            {
                "skill": "stakeholder-analysis",
                "name": "Phân tích Stakeholder",
                "use_case": "Xác định và phân tích tất cả các bên liên quan trong dự án",
                "prompt": (
                    "/stakeholder-analysis\n"
                    "Phân tích Stakeholder cho dự án [Tên dự án]:\n"
                    "- Mô tả ngắn về dự án: [1-3 câu]\n"
                    "- Các phòng ban liên quan: [IT / Kinh doanh / Vận hành / ...]\n\n"
                    "Yêu cầu:\n"
                    "1. Lập danh sách đầy đủ Stakeholders\n"
                    "2. Ma trận Influence vs Interest (Quyền lực vs Quan tâm)\n"
                    "3. Chiến lược quản lý từng nhóm Stakeholder\n"
                    "4. Kế hoạch Communication (Ai nhận thông tin gì, tần suất, kênh nào)"
                ),
                "note": ""
            },
        ]
    },
    # ============================================================
    # NHÓM 8: AGILE / SCRUM
    # ============================================================
    {
        "group": "08 · Agile / Scrum",
        "skills": [
            {
                "skill": "backlog-grooming",
                "name": "Backlog Grooming",
                "use_case": "Xử lý và tinh lọc Product Backlog trước Sprint Planning",
                "prompt": (
                    "/backlog-grooming\n"
                    "Hỗ trợ Backlog Grooming cho Sprint tới của dự án [Tên dự án].\n\n"
                    "Danh sách Backlog Items hiện tại:\n"
                    "[Dán danh sách User Stories / Tasks cần grooming]\n\n"
                    "Yêu cầu:\n"
                    "1. Đánh giá độ rõ ràng (Clarity) và khả năng ước lượng (Estimable) của từng item\n"
                    "2. Phân loại INVEST: Independent / Negotiable / Valuable / Estimable / Small / Testable\n"
                    "3. Đề xuất điểm Story Point (Fibonacci: 1, 2, 3, 5, 8, 13)\n"
                    "4. Tách các item quá lớn (Epic) thành Stories nhỏ hơn\n"
                    "5. Sắp xếp thứ tự ưu tiên theo Business Value và Technical Risk"
                ),
                "note": ""
            },
            {
                "skill": "sprint-planning",
                "name": "Sprint Planning",
                "use_case": "Lập kế hoạch Sprint, phân bổ công việc và xác định Sprint Goal",
                "prompt": (
                    "/sprint-planning\n"
                    "Hỗ trợ lập kế hoạch Sprint [Số Sprint] cho dự án [Tên dự án]:\n"
                    "- Velocity Sprint trước: [N Story Points]\n"
                    "- Team size: [N người]\n"
                    "- Thời gian Sprint: [2 tuần]\n"
                    "- Backlog Items dự kiến đưa vào Sprint: [Danh sách]\n\n"
                    "Yêu cầu:\n"
                    "1. Xác định Sprint Goal\n"
                    "2. Phân bổ công việc theo ngày (Daily Capacity)\n"
                    "3. Kiểm tra Dependencies giữa các Tasks\n"
                    "4. Xác định Definition of Done (DoD) cho Sprint này"
                ),
                "note": ""
            },
        ]
    },
]

# ================================================================
# HÀM TẠO FILE DOCX
# ================================================================
def set_cell_bg(cell, hex_color):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val)
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color_hex="1F3864"):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
    r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return para

def build_docx():
    doc = Document()
    
    # === PAGE MARGINS ===
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === TITLE PAGE ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("BA SKILL LIBRARY\nBỘ SỔ TAY PROMPT MẪU")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    title_para.paragraph_format.space_before = Pt(20)
    title_para.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Tổng hợp toàn bộ Template Prompt tối ưu cho từng kỹ năng nghiệp vụ (BA Skill)\nPhiên bản 1.0 – Dành cho Claude / Grok / ChatGPT")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x5A, 0x5A, 0x8A)
    subtitle.paragraph_format.space_after = Pt(20)

    doc.add_paragraph()  # spacer

    # === HƯỚNG DẪN SỬ DỤNG ===
    add_heading(doc, "Hướng dẫn sử dụng", 1)
    guide_text = (
        "Tài liệu này là Sổ tay Prompt (Prompt Handbook) dành cho BA khi làm việc với AI (Claude, Grok, ChatGPT). "
        "Mỗi dòng trong bảng là một Template Prompt đã được tối ưu hóa, sẵn sàng copy-paste vào giao diện AI.\n\n"
        "Nguyên tắc sử dụng hiệu quả:\n"
        "1. Gọi đúng Skill (VD: /activity-diagram) ở dòng đầu tiên để kích hoạt kỹ năng tương ứng.\n"
        "2. Thay thế các đoạn trong [ngoặc vuông] bằng thông tin thực tế của dự án bạn đang làm.\n"
        "3. Dùng Skill Chaining: Gọi các skill nhỏ trước → Cuối cùng gọi /write-srs-ultra để tổng hợp.\n"
        "4. Với yêu cầu cụ thể hơn, xem thêm folder Claude-Skills hoặc Grok-Skills trong thư viện.\n"
    )
    p = doc.add_paragraph(guide_text)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(14)

    doc.add_page_break()

    # === CÁC NHÓM SKILL ===
    for group_data in PROMPT_TEMPLATES:
        add_heading(doc, group_data["group"], 1)

        for skill in group_data["skills"]:
            # Skill sub-heading
            add_heading(doc, f"► {skill['name']}", 2, color_hex="2E4057")

            # Bảng chi tiết
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Column widths
            table.columns[0].width = Cm(4)
            table.columns[1].width = Cm(13)

            # Header row
            hdr_cells = table.rows[0].cells
            set_cell_bg(hdr_cells[0], "2E4057")
            set_cell_bg(hdr_cells[1], "2E4057")
            hdr_cells[0].merge(hdr_cells[1])
            hdr_para = hdr_cells[0].paragraphs[0]
            hdr_run = hdr_para.add_run(f"🎯  {skill['skill'].upper()}")
            hdr_run.bold = True
            hdr_run.font.size = Pt(10)
            hdr_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            def add_row(label, value, label_bg="DCE6F1", value_bg="F7F9FC", bold_label=True, mono=False):
                row = table.add_row()
                cells = row.cells
                cells[0].width = Cm(4)
                cells[1].width = Cm(13)
                set_cell_bg(cells[0], label_bg)
                set_cell_bg(cells[1], value_bg)
                p0 = cells[0].paragraphs[0]
                r0 = p0.add_run(label)
                r0.bold = bold_label
                r0.font.size = Pt(9.5)
                r0.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                p1 = cells[1].paragraphs[0]
                r1 = p1.add_run(value)
                r1.font.size = Pt(9)
                if mono:
                    r1.font.name = "Consolas"
                    r1.font.color.rgb = RGBColor(0x00, 0x40, 0x00)
                    set_cell_bg(cells[1], "F0F4F0")
                cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

            add_row("📋 Mục đích sử dụng", skill["use_case"])
            add_row("📝 Template Prompt\n(Copy & Paste)", skill["prompt"], mono=True)
            if skill.get("note"):
                add_row("💡 Lưu ý / Mẹo", skill["note"], label_bg="FFF4CE", value_bg="FFFDE7")

            doc.add_paragraph()  # spacer

    # === GHI CHÚ CUỐI ===
    doc.add_page_break()
    add_heading(doc, "Quy tắc Skill Chaining (Kết hợp nhiều Skill)", 1)
    chain_text = (
        "Để đạt kết quả tốt nhất, hãy kết hợp các skill theo chuỗi (Skill Chaining) "
        "trong cùng 1 luồng chat. Dưới đây là 2 kịch bản mẫu thực chiến:\n\n"
        "KỊCH BẢN 1: Phân tích module MES từ đầu\n"
        "  Bước 1: /phan-tich-production-order → Hiểu quy trình\n"
        "  Bước 2: /activity-diagram → Vẽ luồng thao tác tại máy\n"
        "  Bước 3: /phan-tich-business-rules → Bóc tách toàn bộ Business Rules\n"
        "  Bước 4: /thiet-ke-erd → Thiết kế cấu trúc dữ liệu\n"
        "  Bước 5: /phan-tich-mobile-app → Thiết kế màn hình Tablet cho công nhân\n"
        "  Bước 6: /write-srs-ultra → Tổng hợp thành tài liệu SRS hoàn chỉnh\n\n"
        "KỊCH BẢN 2: Chỉ cần viết SRS nhanh\n"
        "  Bước 1: /phan-tich-use-case → Đặc tả Use Case\n"
        "  Bước 2: /phan-tich-business-rules → Bóc tách Business Rules\n"
        "  Bước 3: /write-srs-ultra → Tổng hợp SRS\n\n"
        "LƯU Ý: Trong trường hợp muốn xử lý một yêu cầu cụ thể hơn,\n"
        "có thể kết hợp thêm các skill trong folder Claude-Skills / Grok-Skills của thư viện."
    )
    p = doc.add_paragraph(chain_text)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.name = "Consolas"
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x00)

    doc.save(OUTPUT_FILE)
    print(f"[OK] File DOCX da duoc tao tai: {OUTPUT_FILE}")

if __name__ == "__main__":
    build_docx()
